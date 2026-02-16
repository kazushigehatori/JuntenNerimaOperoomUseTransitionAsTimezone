"""DOCX v3.0 生成スクリプト"""
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(SCRIPT_DIR, "計算集計方法とグラフ表示シート作成手順.docx")

doc = Document()

# スタイル設定
style = doc.styles['Normal']
font = style.font
font.name = 'メイリオ'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

def add_note(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    return p

# ===== タイトル =====
title = doc.add_heading('時間帯別稼働推移\n計算集計方法とグラフ表示シート作成手順', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

p = doc.add_paragraph('バージョン: 3.0 ｜ 最終更新: 2026年2月 ｜ 対象: 手術室管理部門')
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ===== 目次 =====
add_heading('目次', 2)
toc_items = [
    '第1部：計算集計方法',
    '  1.1 ツール概要',
    '  1.2 入力データ仕様',
    '  1.3 集計対象の定義',
    '  1.4 計測区間方式による計数方法',
    '  1.5 曜日別集計',
    '  1.6 計算例',
    '  1.7 出力仕様',
    '第2部：グラフ表示シート作成手順',
    '  2.1 作成の考え方',
    '  2.2 作成手順（Step 1〜9）',
    '  2.3 動作確認',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

# ===== 第1部 =====
add_heading('第1部：計算集計方法', 1)

add_heading('1.1 ツール概要', 2)
doc.add_paragraph(
    '本ツールは、手術実施データから「時間帯ごとに平均何室の手術室が稼働しているか」を集計します。'
    '8:00〜20:00の間、30分おきにスナップショット時刻を設定し、各時刻を中心とした30分間の計測区間'
    '（-14分〜+15分）を定義します。計測区間と手術時間の重なりで使用中の手術室を判定し、ウェイト付き'
    'でカウントして月間の日平均を算出します。また、曜日別（月・火・水・木・金・土）の集計も同時に行い、'
    '各曜日の稼働推移を個別に把握できます。'
)

add_table(
    ['項目', '内容'],
    [
        ['アプリケーション', '時間帯別稼働推移.exe（ダブルクリックで実行）'],
        ['入力ファイル', '時間帯別稼働推移元データ.xlsx'],
        ['出力ファイル', '時間帯別稼働推移-結果.xlsx'],
        ['動作環境', 'Windows 10 / 11（Python不要）'],
    ]
)

add_heading('1.2 入力データ仕様', 2)
doc.add_paragraph('シート「時間帯別稼働推移元データ」のカラム：')
add_table(
    ['カラム名', '型', '説明', '例'],
    [
        ['手術実施管理番号', '文字列/数値', '一意の管理番号', '1000104765'],
        ['手術実施日', '文字列', 'YYYY/MM/DD形式', '2025/09/01'],
        ['曜日', '文字列', '○曜日 形式', '月曜日'],
        ['実施手術室名', '文字列', '手術室番号', '01A, 09, 10'],
        ['入室時刻', '時刻', 'HH:MM:SS形式', '08:44:00'],
        ['麻酔終了時刻', '時刻', 'HH:MM:SS形式', '13:39:00'],
        ['実施申込区分', '文字列', '定時 / 臨時 / 緊急', '定時'],
    ]
)

add_heading('1.3 集計対象の定義', 2)
doc.add_paragraph('対象手術室とウェイト：').runs[0].bold = True
add_table(
    ['手術室', 'ウェイト', '備考'],
    [
        ['01A', '0.5', '小手術室'],
        ['01B', '0.5', '小手術室'],
        ['02, 03, 05, 06, 07, 08, 09, 10', '各1.0', '標準手術室'],
    ]
)
add_note('※ ｱﾝｷﾞｵは本ツールの対象外です（入力データに含まれていても無視されます）。')

doc.add_paragraph('除外曜日：').runs[0].bold = True
add_table(
    ['除外対象', '理由'],
    [
        ['土曜日', '定時帯と稼働パターンが異なるため'],
        ['日曜日', '原則手術なし'],
    ]
)

doc.add_paragraph('集計区分（2パターン）：').runs[0].bold = True
add_table(
    ['区分名', '対象となる実施申込区分', '出力先'],
    [
        ['全手術（緊急含む）', '定時・臨時・緊急', '計算結果シート Row2（B2:Z2）'],
        ['予定手術のみ', '定時 のみ', '計算結果シート Row3（B3:Z3）'],
    ]
)

doc.add_paragraph('曜日集計の定義（Row19-20）：').runs[0].bold = True
add_table(
    ['行', '内容'],
    [
        ['Row19', '曜日集計（土曜日）'],
        ['Row20', '土曜日の稼働を計算'],
    ]
)
add_note('この定義により、曜日別集計では土曜日も計算対象に含まれます。')

# 1.4
add_heading('1.4 計測区間方式による計数方法', 2)

doc.add_paragraph('計測区間の定義').runs[0].bold = True
doc.add_paragraph('各スナップショット時刻に対して、前後に幅を持たせた「計測区間」を設定する。')
add_table(
    ['項目', '定義'],
    [
        ['計測区間の開始時間 A', 'スナップショット時刻 − 14分'],
        ['計測区間の終了時間 B', 'スナップショット時刻 + 15分'],
        ['計測区間の長さ', '30分間'],
        ['全体の計測範囲', '7:46〜20:15（隣接区間は隙間なく連続）'],
    ]
)

doc.add_paragraph('計測区間の一覧（全25区間）：').runs[0].bold = True
snapshot_intervals = [
    ['8:00','7:46〜8:15'],['8:30','8:16〜8:45'],['9:00','8:46〜9:15'],['9:30','9:16〜9:45'],
    ['10:00','9:46〜10:15'],['10:30','10:16〜10:45'],['11:00','10:46〜11:15'],['11:30','11:16〜11:45'],
    ['12:00','11:46〜12:15'],['12:30','12:16〜12:45'],['13:00','12:46〜13:15'],['13:30','13:16〜13:45'],
    ['14:00','13:46〜14:15'],['14:30','14:16〜14:45'],['15:00','14:46〜15:15'],['15:30','15:16〜15:45'],
    ['16:00','15:46〜16:15'],['16:30','16:16〜16:45'],['17:00','16:46〜17:15'],['17:30','17:16〜17:45'],
    ['18:00','17:46〜18:15'],['18:30','18:16〜18:45'],['19:00','18:46〜19:15'],['19:30','19:16〜19:45'],
    ['20:00','19:46〜20:15'],
]
add_table(['スナップショット時刻', '計測区間 A〜B'], snapshot_intervals)

doc.add_paragraph('重なり判定ルール').runs[0].bold = True
doc.add_paragraph(
    '計測区間 [A, B] と手術時間 [C, D] の重なりで稼働中を判定する。\n'
    'C = 手術の入室時刻、D = 手術の麻酔終了時刻。\n\n'
    '判定条件：A ≦ D かつ C ≦ B のとき「稼働中」（閉区間同士の重なり判定）。\n'
    '境界が一致する場合（B=C または A=D）も「稼働中」と判定する。'
)
add_table(
    ['判定', '条件'],
    [
        ['稼働中と判定', '計測区間 [A, B] と手術時間 [C, D] に重なりあり（A ≦ D かつ C ≦ B）'],
        ['カウント値', '稼働中の場合、当該手術室のウェイト値（01A/01Bは0.5、その他は1.0）'],
        ['集計方法', '対象日ごとにカウントし、全対象日の合計 ÷ 対象日数 で日平均を算出'],
        ['丸め', '小数第2位まで（例：6.55）'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('各時刻の値 = Σ（全対象日の当該時刻の稼働室数）÷ 対象日数')
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('なぜ計測区間方式を採用したか').runs[0].bold = True
add_note(
    '旧方式（点判定：入室時刻 ≦ スナップショット時刻 ＜ 麻酔終了時刻）では、スナップショット時刻の'
    'ちょうど1分後に入室した手術は「未稼働」と判定されてしまう。実際のデータでは9時台に開始する手術の'
    '多くが9:01〜9:15に入室しているため、旧方式の9:00の値は実態より大幅に低くなっていた。計測区間方式'
    'により、各時間帯の稼働状況をより実態に即して把握できるようになった。'
)

# 1.5 曜日別集計
add_heading('1.5 曜日別集計', 2)
doc.add_paragraph('全体集計（Row2-3）に加え、月曜日〜土曜日の各曜日ごとに同じ計測区間方式で集計を行います。')

doc.add_paragraph('全体集計と曜日別集計の違い：').runs[0].bold = True
add_table(
    ['項目', '全体集計（Row2-3）', '曜日別集計（Row7以降）'],
    [
        ['対象データ', '土日を除外した全平日', '各曜日のデータのみ'],
        ['除外曜日', '定義シートの除外設定に従う（土曜日・日曜日）', '適用しない（土曜日も計算対象）'],
        ['計測区間・判定ロジック', '同一', '同一'],
        ['出力', 'Row2:全手術、Row3:予定のみ', '曜日ごとに全手術・予定のみの2行'],
    ]
)

doc.add_paragraph('対象曜日と出力先：').runs[0].bold = True
add_table(
    ['曜日', '全手術の出力行', '予定手術の出力行'],
    [
        ['月曜日', 'B7:Z7', 'B8:Z8'],
        ['火曜日', 'B12:Z12', 'B13:Z13'],
        ['水曜日', 'B17:Z17', 'B18:Z18'],
        ['木曜日', 'B22:Z22', 'B23:Z23'],
        ['金曜日', 'B27:Z27', 'B28:Z28'],
        ['土曜日', 'B32:Z32', 'B33:Z33'],
    ]
)
add_note('注：曜日テーブルの並び順は月→火→水→木→金→土です（計算結果シートのレイアウトに準拠）。日曜日は原則手術なしのため集計対象外です。')

doc.add_paragraph('土曜日の集計について：').runs[0].bold = True
add_note('定義シートのRow19-20に「曜日集計（土曜日）：土曜日の稼働を計算」が定義されています。全体集計では除外される土曜日も、曜日別集計では計算対象となります。')

# 1.6 計算例
add_heading('1.6 計算例', 2)
doc.add_paragraph(
    '例1：2025/09/01（月曜日）、スナップショット時刻 = 9:00（計測区間 8:46〜9:15）'
).runs[0].bold = True
doc.add_paragraph(
    'この計測区間と手術時間の重なりを判定：\n'
    '・09号室：入室 8:44 〜 麻酔終了 13:39 → A(8:46) ≦ D(13:39) かつ C(8:44) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・06号室：入室 9:07 〜 麻酔終了 10:24 → A(8:46) ≦ D(10:24) かつ C(9:07) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・10号室：入室 8:41 〜 麻酔終了 15:08 → A(8:46) ≦ D(15:08) かつ C(8:41) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・08号室：入室 9:03 〜 麻酔終了 13:01 → A(8:46) ≦ D(13:01) かつ C(9:03) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・05号室：入室 9:07 〜 麻酔終了 12:23 → A(8:46) ≦ D(12:23) かつ C(9:07) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・02号室：入室 9:03 〜 麻酔終了 10:42 → A(8:46) ≦ D(10:42) かつ C(9:03) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・03号室：入室 9:06 〜 麻酔終了 9:32 → A(8:46) ≦ D(9:32) かつ C(9:06) ≦ B(9:15) → 稼働中（×1.0）\n'
    '・01A号室：入室 14:13 〜 麻酔終了 15:50 → C(14:13) ≦ B(9:15)？→ No（未稼働）\n\n'
    'この日の9:00区間の値 = 合計 7.0室'
)

doc.add_paragraph(
    '例2：旧方式との比較'
).runs[0].bold = True
doc.add_paragraph(
    '旧方式（点判定：入室 ≦ 9:00 < 麻酔終了）では、9:01以降に入室した06, 08, 05, 02, 03号室の5室は'
    '「9:00時点でまだ入室していない」ため未稼働と判定され、この日のカウントは2.0室であった。新方式では'
    '計測区間8:46〜9:15で手術時間との重なりを見るため、9:01〜9:07に入室した手術もすべて捕捉し、7.0室となる。'
)

doc.add_paragraph('参考（全20日平均）：').runs[0].bold = True
add_table(
    ['', '旧方式（点判定）', '新方式（計測区間）'],
    [
        ['全手術 9:00', '2.88', '6.55'],
        ['予定手術 9:00', '2.67', '6.05'],
    ]
)

# 1.7 出力仕様
add_heading('1.7 出力仕様', 2)
doc.add_paragraph('全体集計「計算結果」シートのレイアウト（Row1-3）：').runs[0].bold = True
add_table(
    ['', 'A列', 'B列', 'C列', 'D列', '…', 'Z列'],
    [
        ['Row1', '集計結果', '8:00', '8:30', '9:00', '…', '20:00'],
        ['Row2', '全手術（緊急含む）', '0.0', '1.23', '6.55', '…', '0.82'],
        ['Row3', '予定手術のみ', '0.0', '1.07', '6.05', '…', '0.35'],
    ]
)
add_note('Row1（ヘッダ）とA2:A3（行ラベル）は元ファイルに事前設定済み。exeが書き込むのはB2:Z3の数値（50セル）のみです。')

doc.add_paragraph('曜日別集計のレイアウト（Row5〜33）：').runs[0].bold = True
doc.add_paragraph('各曜日テーブルは以下の構造で繰り返されます（例：月曜日）：')
add_table(
    ['行', 'A列', 'B列〜Z列'],
    [
        ['Row5', '月曜日（曜日ラベル）', ''],
        ['Row6', '集計結果（ヘッダ）', '8:00〜20:00'],
        ['Row7', '全手術（緊急含む）', 'exe実行時に書き込み'],
        ['Row8', '予定手術のみ', 'exe実行時に書き込み'],
        ['Row9', '（空行）', ''],
    ]
)
doc.add_paragraph('同様に火曜日（Row10-14）、水曜日（Row15-19）、木曜日（Row20-24）、金曜日（Row25-29）、土曜日（Row30-33）が続きます。')
add_note('曜日ラベル行、ヘッダ行、行ラベルは元ファイルに事前設定済み。exeが書き込むのは各曜日の値エリア（全手術行と予定手術行のB〜Z列）のみです。')

# ===== 第2部 =====
add_heading('第2部：グラフ表示シート作成手順', 1)

add_heading('2.1 作成の考え方', 2)
doc.add_paragraph(
    '元ファイルに「グラフ表示」シートを追加し、計算結果シートのB2:Z3を参照するグラフをExcelの機能で'
    '作成します。一度作成すれば、以降はexe実行のたびにグラフが自動更新されます。'
)
add_note(
    'なぜExcelでグラフを作るのか：openpyxl（Pythonライブラリ）で面グラフを作成すると、垂直グリッド線の'
    '位置がカテゴリラベルの間に入ってしまう制約があります。Excelの「日付軸」設定を使えばグリッド線を'
    'ラベル位置に正確に配置できるため、グラフはExcelで事前作成する方式を採用しました。'
)

add_heading('2.2 作成手順（Step 1〜9）', 2)

steps = [
    ('Step 1：シート追加',
     '元ファイル（時間帯別稼働推移元データ.xlsx）をExcelで開きます。シートタブを右クリック →「挿入」→ ワークシートを追加し、シート名を「グラフ表示」に変更します。'),
    ('Step 2：グラフ挿入',
     '「グラフ表示」シートを選択した状態で：「挿入」タブ →「グラフ」→「面」→「2-D 面」を選択します。空のグラフが挿入されます。'),
]

for title_text, body in steps:
    doc.add_paragraph(title_text).runs[0].bold = True
    doc.add_paragraph(body)

doc.add_paragraph('Step 3：データ系列1 — 全手術（面グラフ）').runs[0].bold = True
doc.add_paragraph('グラフを右クリック →「データの選択」→「追加」：')
add_table(['項目', '入力値'], [['系列名', '=計算結果!$A$2'], ['系列値', '=計算結果!$B$2:$Z$2']])
doc.add_paragraph('「OK」をクリック。次に「横（項目）軸ラベル」の「編集」をクリック：')
add_table(['項目', '入力値'], [['軸ラベルの範囲', '=計算結果!$B$1:$Z$1']])

doc.add_paragraph('Step 4：データ系列2 — 予定手術のみ（追加して折れ線に変更）').runs[0].bold = True
doc.add_paragraph('グラフを右クリック →「データの選択」→「追加」：')
add_table(['項目', '入力値'], [['系列名', '=計算結果!$A$3'], ['系列値', '=計算結果!$B$3:$Z$3']])
doc.add_paragraph('「OK」→「OK」で閉じます。次にグラフ内の「予定手術のみ」の面部分をクリック → 右クリック →「系列グラフの種類の変更」→「予定手術のみ」を「折れ線」に変更 →「OK」')

doc.add_paragraph('Step 5：色・書式の設定').runs[0].bold = True
doc.add_paragraph('■ 全手術（面グラフ）：面部分をクリック → 右クリック →「データ系列の書式設定」')
add_table(['設定項目', '値'], [['塗りつぶし色', '水色 #BDD7EE'], ['透明度', '30〜40%'], ['枠線', 'なし']])
doc.add_paragraph('■ 予定手術のみ（折れ線）：線をクリック → 右クリック →「データ系列の書式設定」')
add_table(['設定項目', '値'], [['線の色', '濃い青 #2E75B6'], ['線の幅', '2pt'], ['スムージング', 'チェックを入れる']])

doc.add_paragraph('Step 6：縦軸（Y軸）の設定').runs[0].bold = True
doc.add_paragraph('縦軸の数値をダブルクリック →「軸の書式設定」')
add_table(['設定項目', '値'], [['最小値', '0'], ['最大値', '9'], ['目盛間隔（主）', '1']])
doc.add_paragraph('軸タイトルを追加：「（部屋数）」')

doc.add_paragraph('Step 7：横軸（X軸）の設定 ★重要').runs[0].bold = True
doc.add_paragraph('横軸のラベルをダブルクリック →「軸の書式設定」')
add_table(['設定項目', '値'], [['軸の種類', '「日付軸」に変更（これで垂直グリッド線がラベル位置に揃います）'], ['目盛の間隔', '1']])
add_note('ポイント：「テキスト軸」のままだと垂直グリッド線がラベルの間に入ります。必ず「日付軸」に変更してください。')

doc.add_paragraph('Step 8：グリッド線の設定').runs[0].bold = True
doc.add_paragraph('グラフ内をクリック → 右上の「＋」ボタン（グラフ要素）→「目盛線」')
add_table(['設定項目', '操作'], [['水平グリッド線', '「第1主横軸目盛線」にチェック'], ['垂直グリッド線', '「第1主縦軸目盛線」にチェック'], ['グリッド線の色', '各グリッド線をクリック → 薄いグレーに変更']])

doc.add_paragraph('Step 9：凡例・仕上げ・保存').runs[0].bold = True
doc.add_paragraph('凡例はグラフ上部に配置。グラフタイトルは削除または「＜申込区分別＞」等に設定。グラフサイズを適宜調整します。完了したら元ファイル（時間帯別稼働推移元データ.xlsx）を上書き保存します。')

add_heading('2.3 動作確認', 2)
add_note(
    '確認方法：exe を実行して「時間帯別稼働推移-結果.xlsx」を開き、「グラフ表示」シートにグラフが正しく'
    '表示されることを確認してください。計算結果シートのB2:Z3に値が書き込まれ、グラフが自動で反映されます。'
)
add_note(
    '注意事項：\n'
    '・「グラフ表示」シートの名前は変更しないでください\n'
    '・「計算結果」シートのB1:Z1（時刻ヘッダ）やA2:A3（行ラベル）は変更しないでください\n'
    '・exeは「計算結果」シートのB2:Z3および曜日別集計エリア（Row7〜33の値行）に値を書き込みます。'
    '「グラフ表示」シートには一切触れないため、一度グラフを作成すれば毎月のデータ更新はexe実行だけで完了します'
)

# フッター
doc.add_paragraph()
p = doc.add_paragraph('時間帯別稼働推移 — 計算集計方法とグラフ表示シート作成手順 v3.0 (c) 2025')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(OUTPUT)
print(f"DOCX保存完了: {OUTPUT}")
